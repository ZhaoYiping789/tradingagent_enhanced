"""
Word文档报告生成器
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.pyplot as plt
from datetime import datetime
import base64
from io import BytesIO


class TradingReportGenerator:
    """交易分析报告生成器"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """设置文档样式"""
        # 标题样式
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Inches(0.2)
        title_style.font.bold = True
        title_style.font.name = 'Arial'
        
        # 副标题样式
        subtitle_style = self.doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.size = Inches(0.15)
        subtitle_style.font.bold = True
        subtitle_style.font.name = 'Arial'
    
    def add_title(self, title, subtitle=None):
        """添加标题"""
        title_para = self.doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if subtitle:
            subtitle_para = self.doc.add_paragraph(subtitle)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_section(self, title, content):
        """添加章节"""
        self.doc.add_heading(title, 1)
        if isinstance(content, str):
            self.doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                self.doc.add_paragraph(f"• {item}")
    
    def add_table(self, title, data, headers):
        """添加表格"""
        self.doc.add_heading(title, 2)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # 添加数据
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
    
    def add_image(self, image_path, title=None, width=None):
        """添加图片"""
        if title:
            self.doc.add_heading(title, 2)
        
        if os.path.exists(image_path):
            if width:
                self.doc.add_picture(image_path, width=Inches(width))
            else:
                self.doc.add_picture(image_path, width=Inches(6))
        else:
            self.doc.add_paragraph(f"图片未找到: {image_path}")
    
    def add_key_metrics_table(self, metrics):
        """添加关键指标表格"""
        self.doc.add_heading("关键指标", 2)
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "指标"
        hdr_cells[1].text = "数值"
        hdr_cells[2].text = "说明"
        
        # 添加指标
        for metric_name, metric_data in metrics.items():
            row_cells = table.add_row().cells
            row_cells[0].text = metric_name
            row_cells[1].text = str(metric_data.get('value', 'N/A'))
            row_cells[2].text = metric_data.get('description', '')
    
    def save(self, filepath):
        """保存文档"""
        self.doc.save(filepath)
        return filepath


def generate_comprehensive_word_report(
    ticker, 
    current_date,
    market_report=None,
    news_report=None,
    sentiment_report=None,
    fundamentals_report=None,
    quantitative_report=None,
    final_decision=None,
    chart_paths=None
):
    """生成综合Word报告"""
    
    generator = TradingReportGenerator()
    
    # 添加标题页
    generator.add_title(
        f"{ticker} 综合交易分析报告",
        f"分析日期: {current_date}"
    )
    
    # 添加执行摘要
    if final_decision:
        generator.add_section("执行摘要", f"最终交易建议: {final_decision}")
    
    # 添加技术分析
    if market_report:
        generator.add_section("技术分析", market_report)
    
    # 添加图表
    if chart_paths:
        # Handle both dict and list formats for chart_paths
        if isinstance(chart_paths, dict):
            for chart_name, chart_path in chart_paths.items():
                generator.add_image(chart_path, f"{chart_name}图表", width=6)
        elif isinstance(chart_paths, (list, tuple)):
            for idx, chart_path in enumerate(chart_paths, 1):
                generator.add_image(chart_path, f"Technical Analysis Chart {idx}", width=6)
    
    # 添加新闻分析
    if news_report:
        generator.add_section("新闻分析", news_report)
    
    # 添加情感分析
    if sentiment_report:
        generator.add_section("市场情感分析", sentiment_report)
    
    # 添加基本面分析
    if fundamentals_report:
        generator.add_section("基本面分析", fundamentals_report)
    
    # 添加量化分析
    if quantitative_report:
        generator.add_section("量化分析", quantitative_report)
    
    # 保存文档
    results_dir = Path(f"results/{ticker}/{current_date}")
    results_dir.mkdir(parents=True, exist_ok=True)
    
    doc_path = results_dir / f"{ticker}_comprehensive_report_{current_date}.docx"
    generator.save(str(doc_path))
    
    return str(doc_path)


def create_enhanced_summary_chart(df, ticker, current_date, save_path=None):
    """创建增强的汇总图表，包含更多指标"""
    
    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))
    fig.suptitle(f'{ticker} - 综合分析汇总 - {current_date}', fontsize=16, fontweight='bold')
    
    # 图1: 价格和移动平均线 (包含150日和200日)
    ax1.plot(df.index, df['Close'], label='收盘价', linewidth=2, color='black')
    ax1.plot(df.index, df['SMA_10'], label='SMA 10', alpha=0.7, color='blue')
    ax1.plot(df.index, df['SMA_30'], label='SMA 30', alpha=0.7, color='red')
    if 'SMA_150' in df.columns:
        ax1.plot(df.index, df['SMA_150'], label='SMA 150', alpha=0.7, color='green')
    if 'SMA_200' in df.columns:
        ax1.plot(df.index, df['SMA_200'], label='SMA 200', alpha=0.7, color='purple')
    
    # 添加52周高低点
    if '52W_High' in df.columns and '52W_Low' in df.columns:
        ax1.fill_between(df.index, df['52W_Low'], df['52W_High'], alpha=0.1, color='gray', label='52周区间')
    
    ax1.set_title('价格走势与移动平均线')
    ax1.set_ylabel('价格 ($)')
    ax1.legend()
    ax1.grid(True, alpha=0.3)
    
    # 图2: RSI和超买超卖区域
    ax2.plot(df.index, df['RSI'], label='RSI', color='purple', linewidth=2)
    ax2.axhline(y=70, color='r', linestyle='--', alpha=0.7, label='超买线 (70)')
    ax2.axhline(y=30, color='g', linestyle='--', alpha=0.7, label='超卖线 (30)')
    ax2.axhline(y=50, color='gray', linestyle='-', alpha=0.5, label='中性线')
    ax2.fill_between(df.index, 70, 100, alpha=0.2, color='red', label='超买区域')
    ax2.fill_between(df.index, 0, 30, alpha=0.2, color='green', label='超卖区域')
    ax2.set_title('相对强弱指数 (RSI)')
    ax2.set_ylabel('RSI')
    ax2.set_ylim(0, 100)
    ax2.legend()
    ax2.grid(True, alpha=0.3)
    
    # 图3: 成交量
    ax3.bar(df.index, df['Volume']/1e6, alpha=0.6, color='lightblue', label='成交量 (百万)')
    if 'Volume_MA' in df.columns:
        ax3.plot(df.index, df['Volume_MA']/1e6, color='red', linewidth=2, label='成交量均线')
    ax3.set_title('成交量分析')
    ax3.set_ylabel('成交量 (百万)')
    ax3.legend()
    ax3.grid(True, alpha=0.3)
    
    # 图4: 波动率
    ax4.plot(df.index, df['Volatility']*100, color='red', linewidth=2, label='波动率 (%)')
    volatility_ma = df['Volatility'].rolling(20).mean() * 100
    ax4.plot(df.index, volatility_ma, color='blue', linewidth=1, alpha=0.7, label='波动率均线')
    ax4.set_title('价格波动率')
    ax4.set_ylabel('波动率 (%)')
    ax4.legend()
    ax4.grid(True, alpha=0.3)
    
    # 格式化x轴
    for ax in [ax1, ax2, ax3, ax4]:
        ax.tick_params(axis='x', rotation=45)
    
    plt.tight_layout()
    
    # 保存图表
    if save_path:
        plt.savefig(save_path, dpi=300, bbox_inches='tight')
        plt.close()
        return save_path
    else:
        results_dir = Path(f"results/{ticker}/{current_date}")
        results_dir.mkdir(parents=True, exist_ok=True)
        chart_path = results_dir / f"{ticker}_enhanced_summary_{current_date}.png"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        return str(chart_path)
